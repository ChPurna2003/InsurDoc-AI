import os
from io import BytesIO

# ---- Load environment FIRST ----
from dotenv import load_dotenv
load_dotenv(dotenv_path=".env")  # IMPORTANT

# ---- Streamlit Setup ----
import streamlit as st
st.set_page_config(page_title="Insurance GLR App", page_icon="🧾")

# Debug (REMOVE LATER): show whether API key loaded
st.write("DEBUG – Loaded KEY:", os.getenv("OPENROUTER_API_KEY"))

# ---- Other imports ----
from openai import OpenAI
from docx import Document
from pypdf import PdfReader
import json


# ============ LLM Client Setup ============
@st.cache_resource
def get_llm_client():
    """Create and cache OpenRouter LLM client."""
    api_key = os.getenv("OPENROUTER_API_KEY")
    if not api_key:
        raise ValueError("OPENROUTER_API_KEY not found in .env")

    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=api_key,
        default_headers={
            "Authorization": f"Bearer {api_key}",
            "HTTP-Referer": "https://openrouter.ai",
            "X-Title": "Insurance-GLR-Streamlit-App"
        }
    )
    return client


# ============ Helpers ============
def extract_text_from_pdfs(pdf_files):
    """Extract text from all PDF photo reports."""
    text = []
    for pdf in pdf_files:
        reader = PdfReader(pdf)
        for page in reader.pages:
            text.append(page.extract_text() or "")
    return "\n\n".join(text)


def extract_text_from_docx(docfile):
    """Extract text from DOCX template (paragraphs + tables)."""
    doc = Document(docfile)
    lines = []

    for p in doc.paragraphs:
        lines.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    lines.append(p.text)

    return "\n".join(lines)


def call_llm_to_get_field_mapping(template_text, report_text):
    """Ask LLM (DeepSeek Free) to infer the template key-value pairs."""
    client = get_llm_client()

    system_prompt = """
You extract structured data from insurance photo reports to fill a DOCX template.

Return JSON:
{
  "fields": {
     "<template_placeholder>": "<value>",
     ...
  }
}

If a value is missing, return "".
Only return valid JSON. No explanation text.
"""

    user_prompt = f"""
=== TEMPLATE TEXT ===
{template_text}

=== REPORT TEXT ===
{report_text}
"""

    response = client.chat.completions.create(
    model="mistralai/mistral-7b-instruct:free",
    messages=[
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ],
    response_format={"type": "json_object"},
    temperature=0.2,
)


    raw = response.choices[0].message.content
    data = json.loads(raw)
    return data.get("fields", {})


def fill_template_docx(template_file, field_values):
    """Replace placeholders inside DOCX."""
    doc = Document(template_file)

    def replace(text):
        for k, v in field_values.items():
            text = text.replace(k, v)
        return text

    for p in doc.paragraphs:
        p.text = replace(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace(cell.text)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# ============ Streamlit UI ============

st.title("🧾 Insurance GLR Pipeline (DeepSeek Version)")
st.write("""
Upload:
- A **.docx** insurance template  
- One or more **.pdf** photo reports  

This app will:
1. Extract text from PDF  
2. Use DeepSeek LLM (Free)  
3. Detect template fields  
4. Fill the insurance DOCX  
""")

template_file = st.file_uploader("Upload Template (.docx)", type=["docx"])
pdf_files = st.file_uploader("Upload Photo Reports (.pdf)", type=["pdf"], accept_multiple_files=True)

if st.button("Generate Filled Template"):
    if not template_file:
        st.error("Please upload a DOCX template.")
        st.stop()

    if not pdf_files:
        st.error("Please upload at least one PDF photo report.")
        st.stop()

    with st.spinner("Extracting PDF text..."):
        reports_text = extract_text_from_pdfs(pdf_files)

    with st.spinner("Reading template..."):
        temp_bytes = template_file.read()
        template_text = extract_text_from_docx(BytesIO(temp_bytes))

    with st.spinner("Calling LLM (DeepSeek Free)..."):
        field_mapping = call_llm_to_get_field_mapping(template_text, reports_text)

    st.subheader("Detected Field Mapping")
    st.json(field_mapping)

    with st.spinner("Generating DOCX..."):
        result = fill_template_docx(BytesIO(temp_bytes), field_mapping)

    st.success("Completed! Download your filled template:")
    st.download_button(
        "⬇ Download Filled DOCX",
        data=result,
        file_name="filled_template.docx"
    )
